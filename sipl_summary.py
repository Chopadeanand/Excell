"""
sipl_summary.py  —  HiRATE Report → Word Summary Document generator
====================================================================
Reads one or more *_REPORT.xlsx files (output of sipl_report.py),
extracts data for each project, and builds HiRATE_Summary.docx with:

  - Cover / Executive Summary (overall KPIs across all projects)
  - Per-project section with:
      - KPI scorecard table
      - Category-wise observations table (with Risk Band)
      - Division-wise risk ranking table (sorted by % Issues ↓)

Dependencies:  pip install openpyxl python-docx
Usage:
  python sipl_summary.py                        # reads *_REPORT.xlsx in same folder
  python sipl_summary.py file1.xlsx file2.xlsx  # specific files
  python sipl_summary.py --output MySummary.docx
"""

import os, sys, glob, tempfile, argparse, re
import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

DEFAULT_OUTPUT = "HiRATE_Summary.docx"
SCRIPT_DIR     = os.path.dirname(os.path.abspath(__file__))

# ── Colour constants (RGB) ────────────────────────────────────────────────────
NAVY        = RGBColor(0x1F, 0x4E, 0x79)
BLUE_MID    = RGBColor(0x2E, 0x75, 0xB6)
GREEN       = RGBColor(0x27, 0xAE, 0x60)
RED         = RGBColor(0xC0, 0x39, 0x2B)
ORANGE      = RGBColor(0xED, 0x7D, 0x31)
YELLOW      = RGBColor(0xFF, 0xC0, 0x00)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_LIGHT  = RGBColor(0xF2, 0xF2, 0xF2)
GRAY_MID    = RGBColor(0xD9, 0xD9, 0xD9)
GRAY_DARK   = RGBColor(0x59, 0x59, 0x59)

# Hex strings for table shading
HEX_NAVY    = "1F4E79"
HEX_BLUE    = "2E75B6"
HEX_GREEN   = "27AE60"
HEX_RED     = "C0392B"
HEX_ORANGE  = "ED7D31"
HEX_YELLOW  = "FFC000"
HEX_LGRAY   = "F2F2F2"
HEX_MGRAY   = "D9D9D9"
HEX_WHITE   = "FFFFFF"


# ══════════════════════════════════════════════════════════════════════════════
# DATA EXTRACTION  (mirrors sipl_dashboard.py / sipl_ppt.py)
# ══════════════════════════════════════════════════════════════════════════════

def extract_report_data(xlsx_path):
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb["Sheet1"]

    # Project name from second sheet  e.g. ALL_MKTPL_Ratings_List → MKTPL
    data_sheet = next((s for s in wb.sheetnames if s != "Sheet1"), None)
    proj_name  = "Project"
    if data_sheet:
        m = re.match(r"ALL_(.+?)_Ratings_List", data_sheet)
        proj_name = m.group(1) if m else data_sheet

    # Category rows 3-9, cols 10(label) 11(total) 12(issues) 13(pct)
    cat_labels, cat_total, cat_issues, cat_pct = [], [], [], []
    for r in range(3, 10):
        lbl = ws.cell(r, 10).value
        if lbl is None:
            continue
        if str(lbl).lower().startswith("overall"):
            continue
        cat_labels.append(str(lbl))
        cat_total.append(int(ws.cell(r, 11).value or 0))
        cat_issues.append(int(ws.cell(r, 12).value or 0))
        pv = ws.cell(r, 13).value
        cat_pct.append(round(float(pv) * 100, 2) if isinstance(pv, float) else 0.0)

    # Overall KPIs: row 3, cols 16(total) 17(satisfactory) 18(observations)
    total_kpi = int(ws.cell(3, 16).value or 0)
    sat_kpi   = int(ws.cell(3, 17).value or 0)
    obs_kpi   = int(ws.cell(3, 18).value or 0)
    sat_pct   = round(sat_kpi / total_kpi * 100, 1) if total_kpi else 0
    obs_pct   = round(obs_kpi / total_kpi * 100, 1) if total_kpi else 0

    # Division rows 201+, cols 1(label) 2(total) 3(issues) 4(pct)
    div_labels, div_total, div_issues, div_pct = [], [], [], []
    r = 201
    while ws.cell(r, 1).value:
        div_labels.append(str(ws.cell(r, 1).value))
        div_total.append(int(ws.cell(r, 2).value or 0))
        div_issues.append(int(ws.cell(r, 3).value or 0))
        pv = ws.cell(r, 4).value
        div_pct.append(round(float(pv) * 100, 2) if isinstance(pv, float) else 0.0)
        r += 1

    # Sort division by % Issues descending
    if div_labels:
        combined = sorted(zip(div_pct, div_labels, div_total, div_issues), reverse=True)
        div_pct, div_labels, div_total, div_issues = map(list, zip(*combined))

    return {
        "proj_name":    proj_name,
        "cat_labels":   cat_labels,
        "cat_total":    cat_total,
        "cat_issues":   cat_issues,
        "cat_pct":      cat_pct,
        "total":        total_kpi,
        "satisfactory": sat_kpi,
        "observations": obs_kpi,
        "sat_pct":      sat_pct,
        "obs_pct":      obs_pct,
        "div_labels":   div_labels,
        "div_total":    div_total,
        "div_issues":   div_issues,
        "div_pct":      div_pct,
    }


# ══════════════════════════════════════════════════════════════════════════════
# DOCX HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _shade_cell(cell, hex_color):
    """Apply solid background shading to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, **kwargs):
    """Set borders on a cell. keys: top, bottom, left, right — each a dict with
    style, size, color."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        if side in kwargs:
            el = OxmlElement(f"w:{side}")
            for k, v in kwargs[side].items():
                el.set(qn(f"w:{k}"), str(v))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def _bold_run(para, text, size_pt=11, color=None, italic=False):
    run = para.add_run(text)
    run.bold  = True
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def _plain_run(para, text, size_pt=11, color=None, italic=False):
    run = para.add_run(text)
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    run.italic = italic
    return run


def _set_para_spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"),  str(after))
    pPr.append(spacing)


def _risk_band(pct):
    """Return (label, hex_color) for a given % of issues."""
    if pct == 0:
        return "Clean",    HEX_GREEN
    elif pct < 15:
        return "Low",      HEX_BLUE
    elif pct < 30:
        return "Moderate", HEX_YELLOW
    else:
        return "High",     HEX_RED


def _kpi_table(doc, total, sat, obs, sat_pct, obs_pct):
    """3-column KPI scorecard table."""
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style     = "Table Grid"

    headers = ["TOTAL OBSERVATIONS", "SATISFACTORY", "ISSUES FOUND"]
    colors  = [HEX_NAVY, HEX_GREEN, HEX_RED]
    values  = [f"{total:,}", f"{sat_pct}%", f"{obs_pct}%"]
    subs    = ["Audits Conducted", f"{sat:,} observations", f"{obs:,} issues"]

    row = tbl.rows[0]
    for i, cell in enumerate(row.cells):
        _shade_cell(cell, colors[i])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.width = Inches(2.5)

        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _bold_run(p1, headers[i], size_pt=8, color=WHITE)
        _set_para_spacing(p1, before=80, after=0)

        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _bold_run(p2, values[i], size_pt=22, color=WHITE)
        _set_para_spacing(p2, before=0, after=0)

        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _plain_run(p3, subs[i], size_pt=9, color=WHITE)
        _set_para_spacing(p3, before=0, after=80)

    return tbl


def _cat_table(doc, labels, totals, issues, pcts):
    """Category-wise observations table."""
    hdrs = ["Category", "Total Audited", "Issues", "% Issues", "Risk Band"]
    col_w = [Inches(2.5), Inches(1.2), Inches(1.0), Inches(1.0), Inches(1.0)]

    tbl = doc.add_table(rows=1 + len(labels), cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style     = "Table Grid"

    # Header row
    hrow = tbl.rows[0]
    for i, (cell, hdr) in enumerate(zip(hrow.cells, hdrs)):
        _shade_cell(cell, HEX_NAVY)
        cell.width = col_w[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _bold_run(p, hdr, size_pt=9, color=WHITE)
        _set_para_spacing(p, before=60, after=60)

    # Data rows
    for ri, (lbl, tot, iss, pct) in enumerate(zip(labels, totals, issues, pcts)):
        row  = tbl.rows[ri + 1]
        band, band_hex = _risk_band(pct)
        bg   = HEX_LGRAY if ri % 2 == 0 else HEX_WHITE

        vals = [lbl, f"{tot:,}", str(iss), f"{pct:.1f}%", band]
        for ci, (cell, val) in enumerate(zip(row.cells, vals)):
            cell.width = col_w[ci]
            if ci == 4:          # Risk Band cell gets colour
                _shade_cell(cell, band_hex)
                txt_col = WHITE
            else:
                _shade_cell(cell, bg)
                txt_col = None

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            if ci == 1 and iss > 0:     # Bold if highest issues
                fn = _bold_run if (iss == max(issues)) else _plain_run
            elif ci == 2 and iss > 0:
                fn = _bold_run if (iss == max(issues)) else _plain_run
            else:
                fn = _plain_run
            fn(p, val, size_pt=9, color=WHITE if ci == 4 else txt_col)
            _set_para_spacing(p, before=50, after=50)

    return tbl


def _div_table(doc, labels, totals, issues, pcts):
    """Division-wise risk ranking table, sorted by % issues (already sorted)."""
    hdrs = ["#", "Asset / Division", "Total", "Issues", "% Issues", "Risk"]
    col_w = [Inches(0.35), Inches(2.15), Inches(0.9), Inches(0.9), Inches(1.0), Inches(1.0)]

    tbl = doc.add_table(rows=1 + len(labels), cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style     = "Table Grid"

    # Header
    hrow = tbl.rows[0]
    for i, (cell, hdr) in enumerate(zip(hrow.cells, hdrs)):
        _shade_cell(cell, HEX_BLUE)
        cell.width = col_w[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _bold_run(p, hdr, size_pt=9, color=WHITE)
        _set_para_spacing(p, before=60, after=60)

    # Data
    for ri, (lbl, tot, iss, pct) in enumerate(zip(labels, totals, issues, pcts)):
        row  = tbl.rows[ri + 1]
        band, band_hex = _risk_band(pct)
        bg   = HEX_LGRAY if ri % 2 == 0 else HEX_WHITE

        vals = [str(ri + 1), lbl, f"{tot:,}", str(iss) if iss else "0",
                f"{pct:.1f}%", band]
        for ci, (cell, val) in enumerate(zip(row.cells, vals)):
            cell.width = col_w[ci]
            if ci == 5:
                _shade_cell(cell, band_hex)
                txt_col = WHITE
            else:
                _shade_cell(cell, bg)
                txt_col = None

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT
            fn = _bold_run if (ci == 3 and iss > 0) else _plain_run
            fn(p, val, size_pt=9, color=WHITE if ci == 5 else txt_col)
            _set_para_spacing(p, before=50, after=50)

    return tbl


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT BUILD
# ══════════════════════════════════════════════════════════════════════════════

def build_summary(projects, output_path):
    doc = Document()

    # ── Page setup: A4, narrow margins ────────────────────────────────────────
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    # ── Default style ─────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    month_year = datetime.now().strftime("%B %Y")

    # ══ COVER / EXECUTIVE SUMMARY ═════════════════════════════════════════════

    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _bold_run(p, "HiRATE Audit Report", size_pt=22, color=NAVY)
    _set_para_spacing(p, before=0, after=60)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _plain_run(p, f"Dashboard Summary — {month_year}", size_pt=13, color=GRAY_DARK)
    _set_para_spacing(p, before=0, after=60)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _plain_run(p, "CONFIDENTIAL — For Management Use Only",
               size_pt=10, color=GRAY_DARK, italic=True)
    _set_para_spacing(p, before=0, after=200)

    # Executive summary paragraph
    total_all = sum(d["total"]        for d in projects)
    sat_all   = sum(d["satisfactory"] for d in projects)
    obs_all   = sum(d["observations"] for d in projects)
    n_proj    = len(projects)
    sat_pct_all = round(sat_all / total_all * 100, 1) if total_all else 0
    obs_pct_all = round(obs_all / total_all * 100, 1) if total_all else 0

    p = doc.add_heading("1. Executive Summary", level=1)
    p.runs[0].font.color.rgb = NAVY

    p = doc.add_paragraph()
    _plain_run(p,
        f"This report summarises HiRATE audit findings for {month_year} across "
        f"{n_proj} project(s). Combined total: {total_all:,} observations — "
        f"{sat_all:,} ({sat_pct_all}%) satisfactory, "
        f"{obs_all:,} ({obs_pct_all}%) requiring attention.",
        size_pt=11)
    _set_para_spacing(p, before=0, after=120)

    # Overall KPI scorecard
    p = doc.add_paragraph()
    _bold_run(p, "Overall Compliance Scorecard", size_pt=11, color=NAVY)
    _set_para_spacing(p, before=60, after=80)

    _kpi_table(doc, total_all, sat_all, obs_all, sat_pct_all, obs_pct_all)

    doc.add_paragraph()  # spacer

    # ══ PER-PROJECT SECTIONS ══════════════════════════════════════════════════

    p = doc.add_heading("2. Project-by-Project Analysis", level=1)
    p.runs[0].font.color.rgb = NAVY

    for idx, d in enumerate(projects):
        proj = d["proj_name"]
        num  = f"2.{idx + 1}"

        # --- Section heading ---
        p = doc.add_heading(f"{num} {proj}", level=2)
        p.runs[0].font.color.rgb = BLUE_MID

        # --- Project KPI scorecard ---
        _kpi_table(doc, d["total"], d["satisfactory"], d["observations"],
                   d["sat_pct"], d["obs_pct"])
        doc.add_paragraph()  # spacer

        # --- Narrative summary ---
        # Find highest issue category
        if d["cat_issues"]:
            max_cat_idx = d["cat_issues"].index(max(d["cat_issues"]))
            max_cat_lbl = d["cat_labels"][max_cat_idx]
            max_cat_iss = d["cat_issues"][max_cat_idx]
            max_cat_pct = d["cat_pct"][max_cat_idx]

            min_cat_pct = min(d["cat_pct"])
            min_cat_lbl = d["cat_labels"][d["cat_pct"].index(min_cat_pct)]

            narrative = (
                f"For the {proj} project, {d['total']:,} audit observations were recorded. "
                f"Of these, {d['satisfactory']:,} ({d['sat_pct']}%) were satisfactory and "
                f"{d['observations']:,} ({d['obs_pct']}%) required attention. "
                f"Across {len(d['cat_labels'])} audit categories, "
                f"{max_cat_lbl} had the most issues ({max_cat_iss} issues, "
                f"{max_cat_pct:.1f}% of its total). "
                f"Lowest issue rate: {min_cat_lbl} ({min_cat_pct:.1f}%)."
            )
        else:
            narrative = (
                f"For the {proj} project, {d['total']:,} audit observations were recorded. "
                f"Of these, {d['satisfactory']:,} ({d['sat_pct']}%) were satisfactory and "
                f"{d['observations']:,} ({d['obs_pct']}%) required attention."
            )

        p = doc.add_paragraph()
        _plain_run(p, narrative, size_pt=10)
        _set_para_spacing(p, before=80, after=80)

        # --- Category table ---
        p = doc.add_paragraph()
        _bold_run(p, "Category-wise Observations", size_pt=10, color=NAVY)
        _set_para_spacing(p, before=60, after=60)

        _cat_table(doc, d["cat_labels"], d["cat_total"], d["cat_issues"], d["cat_pct"])
        doc.add_paragraph()

        # --- Division summary narrative ---
        if d["div_labels"]:
            top_div      = d["div_labels"][0]
            top_div_pct  = d["div_pct"][0]
            top_div_iss  = d["div_issues"][0]
            high_risk_n  = sum(1 for p2 in d["div_pct"] if p2 >= 30)
            low_risk_n   = sum(1 for p2 in d["div_pct"] if 0 < p2 < 15)
            zero_divs    = [l for l, i in zip(d["div_labels"], d["div_issues"]) if i == 0]

            div_narrative = (
                f"Division-wise (sorted by % issues): {top_div} had the highest rate "
                f"at {top_div_pct:.1f}% ({top_div_iss} issues). "
            )
            if high_risk_n:
                high_risk_names = [l for l, p2 in zip(d["div_labels"], d["div_pct"]) if p2 >= 30]
                div_narrative += f"High Risk (≥30%): {', '.join(high_risk_names[:5])}. "
            div_narrative += f"{low_risk_n} division(s) in Low Risk (<15%)."
            if zero_divs:
                div_narrative += f" Zero issues in: {', '.join(zero_divs[:5])}."

            p = doc.add_paragraph()
            _plain_run(p, div_narrative, size_pt=10)
            _set_para_spacing(p, before=60, after=60)

        # --- Division table ---
        p = doc.add_paragraph()
        _bold_run(p, "Division-wise Risk Ranking (sorted by % Issues ↓)",
                  size_pt=10, color=NAVY)
        _set_para_spacing(p, before=60, after=60)

        _div_table(doc, d["div_labels"], d["div_total"], d["div_issues"], d["div_pct"])

        # Page break between projects (not after last)
        if idx < len(projects) - 1:
            doc.add_page_break()

    doc.save(output_path)
    print(f"\n✓  Saved: {output_path}")
    print(f"   {len(projects)} project(s) summarised")
    return True


# ══════════════════════════════════════════════════════════════════════════════
# ENTRY POINTS
# ══════════════════════════════════════════════════════════════════════════════

def find_report_files(folder):
    return sorted(glob.glob(os.path.join(folder, "*_REPORT.xlsx")))


def build_summary_from_files(report_files, output_path):
    """Build summary from a list of file paths."""
    projects = []
    for f in report_files:
        print(f"  Reading: {os.path.basename(f)}")
        try:
            d = extract_report_data(f)
            projects.append(d)
            print(f"    → {d['proj_name']}  cats={len(d['cat_labels'])}  divs={len(d['div_labels'])}")
        except Exception as e:
            print(f"    ⚠ Skipped: {e}")

    if not projects:
        print("No valid report files.")
        return False

    return build_summary(projects, output_path)


def generate_summary_from_reports(report_bytes_list, output_path):
    """Called by sipl_app.py. report_bytes_list = [(filename, bytes), ...]"""
    tmp = []
    try:
        for fname, fbytes in report_bytes_list:
            tf = tempfile.NamedTemporaryFile(
                suffix=".xlsx", delete=False,
                prefix=os.path.splitext(fname)[0] + "_"
            )
            tf.write(fbytes)
            tf.close()
            tmp.append(tf.name)
        return build_summary_from_files(tmp, output_path)
    finally:
        for f in tmp:
            try:
                os.unlink(f)
            except Exception:
                pass


def main():
    ap = argparse.ArgumentParser(description="HiRATE Summary DOCX generator")
    ap.add_argument("files",   nargs="*", help="*_REPORT.xlsx files")
    ap.add_argument("--output", "-o", default=None, help="Output .docx path")
    args = ap.parse_args()

    files = [os.path.abspath(f) for f in args.files] if args.files \
            else find_report_files(SCRIPT_DIR)

    if not files:
        print(f"ERROR: No *_REPORT.xlsx files found in {SCRIPT_DIR}")
        sys.exit(1)

    out = os.path.abspath(args.output) if args.output \
          else os.path.join(SCRIPT_DIR, DEFAULT_OUTPUT)

    print(f"\nHiRATE Summary Generator  ·  {len(files)} file(s)  →  {out}\n{'─'*60}")
    sys.exit(0 if build_summary_from_files(files, out) else 1)


if __name__ == "__main__":
    main()
